import streamlit as st
from fpdf import FPDF
import base64
import textwrap
import streamlit.components.v1 as components
from utils import apply_custom_css, render_top_nav
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

st.set_page_config(page_title="Resume Creator", page_icon="📝", layout="wide")
apply_custom_css()
render_top_nav()

st.markdown("# 📝 Smart Resume Creator")
st.markdown("### Build a professional, ATS-friendly resume.")

# --- JOB DESCRIPTIONS DATA ---
JOB_DESCRIPTIONS = {
    "Select a Role": "",
    "Software Engineer": "Highly motivated Software Engineer with experience in designing and developing scalable software solutions. Proficient in algorithms, data structures, and full-stack development. Passionate about writing clean, maintainable code and solving complex technical challenges.",
    "Project Manager": "Results-oriented Project Manager with a proven track record of leading cross-functional teams to deliver projects on time and within budget. Skilled in agile methodologies, risk management, and stakeholder communication.",
    "Data Analyst": "Detail-oriented Data Analyst skilled in collecting, cleaning, and interpreting complex datasets. Proficient in SQL, Python, and data visualization tools to drive business decision-making through actionable insights.",
    "HR Recruiter": "Experienced HR Recruiter with widespread knowledge of talent acquisition, employee relations, and organizational development. Committed to building strong teams and fostering a positive company culture.",
    "Full Stack Developer": "Versatile Full Stack Developer with expertise in both front-end and back-end technologies. capable of building robust web applications from concept to deployment, ensuring high performance and responsiveness.",
    "Machine Learning Engineer": "Innovative Machine Learning Engineer with a strong background in statistical modeling and deep learning. Experienced in building and deploying predictive models to solve real-world problems."
}

# --- SESSION STATE INITIALIZATION ---
if 'education_entries' not in st.session_state:
    st.session_state['education_entries'] = [{'degree': '', 'school': '', 'year': ''}]
if 'internship_entries' not in st.session_state:
    st.session_state['internship_entries'] = [{'role': '', 'company': '', 'dates': '', 'description': ''}]
if 'project_entries' not in st.session_state:
    # Initialize with 2 empty project slots as requested
    st.session_state['project_entries'] = [{'title': '', 'link': '', 'description': ''}, {'title': '', 'link': '', 'description': ''}]
if 'certifications' not in st.session_state:
    st.session_state['certifications'] = ""
if 'achievements' not in st.session_state:
    st.session_state['achievements'] = ""
if 'objective' not in st.session_state:
    st.session_state['objective'] = ""

def add_education():
    st.session_state['education_entries'].append({'degree': '', 'school': '', 'year': ''})

def remove_education(index):
    if len(st.session_state['education_entries']) > 1:
        st.session_state['education_entries'].pop(index)

def add_internship():
    st.session_state['internship_entries'].append({'role': '', 'company': '', 'dates': '', 'description': ''})

def remove_internship(index):
    if len(st.session_state['internship_entries']) > 1:
        st.session_state['internship_entries'].pop(index)

def add_project():
    st.session_state['project_entries'].append({'title': '', 'link': '', 'description': ''})

def remove_project(index):
    if len(st.session_state['project_entries']) > 1:
        st.session_state['project_entries'].pop(index)

# --- PDF GENERATION LOGIC ---
class PDF(FPDF):
    def header(self):
        pass  # We manually handle the header in the body to control exact placement

    def footer(self):
        pass

def create_ats_resume():
    pdf = PDF()
    pdf.add_page()
    
    # Colors
    HEADER_COLOR = (44, 62, 80) # Dark Blue/Grey
    LINK_COLOR = (0, 0, 255) # Blue
    TEXT_COLOR = (0, 0, 0)

    # --- SANITIZATION HELPER ---
    def sanitize_text(text):
        if not text:
            return ""
        # Replace common incompatible characters
        # Using Middle Dot (·) which is chr(183) in Latin-1 as a closer alternative to Bullet (•)
        text = text.replace('\u2022', chr(183)) 
        text = text.replace('\u2013', '-')  # En dash
        text = text.replace('\u2014', '-')  # Em dash
        text = text.replace('\u201c', '"').replace('\u201d', '"') # Smart quotes
        text = text.replace('\u2018', "'").replace('\u2019', "'") # Smart single quotes
        
        # Final fallback: encode to latin-1, replacing errors with '?'
        return text.encode('latin-1', 'replace').decode('latin-1')
    
    # --- NAME & CONTACT ---
    # Reduced font size from 24 to 22 for compactness
    pdf.set_font("Times", "B", 22)
    pdf.set_text_color(*TEXT_COLOR)
    pdf.cell(0, 10, name.upper() if name else "YOUR NAME", ln=True, align="C")
    
    pdf.set_font("Times", "", 10)
    pdf.set_text_color(*LINK_COLOR)
    
    # Construct contact line 1
    contact_info = []
    if phone: contact_info.append(sanitize_text(phone))
    if email: contact_info.append(sanitize_text(email))
    # Reduced line spacing slightly
    pdf.cell(0, 5, "   ".join(contact_info), ln=True, align="C")
    
    # Construct contact line 2 (Links)
    links = []
    if linkedin: links.append("LinkedIn")
    if github: links.append("GitHub")
    if hackerrank: links.append("HackerRank")
    if leetcode: links.append("LeetCode")
    
    # Simple centered text for links
    pdf.cell(0, 5, "   ".join(links), ln=True, align="C")
    
    # Reduced line break from 5 to 2
    pdf.ln(2)

    # Helper for Section Headings
    def section_heading(title):
        pdf.set_font("Times", "B", 12)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(0, 6, title.upper(), ln=True)
        pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 190, pdf.get_y()) # Horizontal line
        # Reduced spacing after line from 2 to 1
        pdf.ln(1)

    # --- OBJECTIVE ---
    if st.session_state['objective']:
        section_heading("Objective")
        pdf.set_font("Times", "", 10) # Reduced from 11 to 10
        pdf.multi_cell(0, 4, sanitize_text(st.session_state['objective'])) # Reduced line height from 5 to 4
        pdf.ln(2) # Reduced from 4 to 2

    # --- EDUCATION ---
    has_education = any(e['degree'] for e in st.session_state['education_entries'])
    if has_education:
        section_heading("Education")
        for edu in st.session_state['education_entries']:
            if edu['degree']:
                pdf.set_font("Times", "B", 10) # Reduced from 11 to 10
                # Cell width 140 for degree, rest for year
                pdf.cell(140, 5, sanitize_text(edu['degree']), ln=0)
                pdf.set_font("Times", "B", 10) 
                pdf.cell(0, 5, sanitize_text(edu['year']), ln=1, align="R")
                
                pdf.set_font("Times", "", 10)
                pdf.cell(0, 5, sanitize_text(edu['school']), ln=1)
                # Removed extra ln(2) between entries to compact list
        pdf.ln(2)

    # --- TECHNICAL SKILLS ---
    has_skills = any([skill_prog_langs, skill_web_dev, skill_core, skill_tools, skill_soft])
    if has_skills:
        section_heading("Technical Skills")
        
        def add_skill_row(label, value):
            if value:
                pdf.set_font("Times", "B", 10) # Reduced from 11 to 10
                
                category_width = 50 # Reduced from 55
                start_x = pdf.get_x()
                start_y = pdf.get_y()
                
                # Print Category Label
                pdf.cell(category_width, 4, label, ln=0) # Reduced height 5->4
                
                # Move to Value position using margin manipulation
                value_start_x = start_x + category_width
                original_l_margin = pdf.l_margin
                pdf.set_left_margin(value_start_x)
                
                # Print Value
                pdf.set_font("Times", "", 10)
                # Use multi_cell but capture height to know where to reset
                pdf.multi_cell(0, 4, sanitize_text(value)) # Reduced height 5->4
                
                # IMPORTANT FIX: Reset left margin AND X position to avoid staircase effect
                pdf.set_left_margin(original_l_margin)
                pdf.set_x(original_l_margin)

        if skill_prog_langs: add_skill_row("Programming Languages", skill_prog_langs)
        if skill_web_dev: add_skill_row("Web Development", skill_web_dev)
        if skill_core: add_skill_row("Core Concepts", skill_core)
        if skill_tools: add_skill_row("Tools/Technologies", skill_tools)
        if skill_soft: add_skill_row("Soft Skills", skill_soft)
            
        pdf.ln(2) # Reduced from 4 to 2

    # --- INTERNSHIPS ---
    has_internships = any(i['role'] for i in st.session_state['internship_entries'])
    if has_internships:
        section_heading("Internships")
        for intern in st.session_state['internship_entries']:
            if intern['role']:
                pdf.set_font("Times", "B", 10) # Reduced 11->10
                pdf.cell(0, 4, sanitize_text(intern['role']), ln=1) # Height 5->4
                
                if intern['company']:
                    pdf.set_font("Times", "I", 10)
                    pdf.cell(140, 4, sanitize_text(intern['company']), ln=0) # Height 5->4
                    pdf.cell(0, 4, sanitize_text(intern['dates']), ln=1, align="R") 
                
                if intern['description']:
                    pdf.set_font("Times", "", 10)
                    lines = intern['description'].split('\n')
                    for line in lines:
                        line = line.strip()
                        if line:
                            sanitized = sanitize_text(line)
                            if not sanitized.strip().startswith(chr(183)) and not sanitized.strip().startswith("-"):
                                sanitized = chr(183) + " " + sanitized
                            pdf.multi_cell(0, 4, sanitized) # Height 5->4
                pdf.ln(2) # Reduced 3->2

    # --- PROJECTS ---
    has_projects = any(p['title'] for p in st.session_state['project_entries'])
    if has_projects:
        section_heading("Projects")
        for proj in st.session_state['project_entries']:
            if proj['title']:
                pdf.set_font("Times", "B", 10)
                pdf.cell(150, 4, sanitize_text(proj['title']), ln=0)
                
                if proj['link']:
                    pdf.set_font("Times", "I", 9) # Smaller for link
                    pdf.set_text_color(*LINK_COLOR)
                    pdf.cell(0, 4, "Project Link", ln=1, align="R", link=sanitize_text(proj['link']))
                    pdf.set_text_color(*TEXT_COLOR)
                else:
                    pdf.ln(4)
                
                if proj['description']:
                    pdf.set_font("Times", "", 10)
                    lines = proj['description'].split('\n')
                    for line in lines:
                        line = line.strip()
                        if line:
                            sanitized = sanitize_text(line)
                            if not sanitized.strip().startswith(chr(183)) and not sanitized.strip().startswith("-"):
                                sanitized = chr(183) + " " + sanitized
                            pdf.multi_cell(0, 4, sanitized)
                pdf.ln(2)

    # --- CERTIFICATIONS ---
    if st.session_state['certifications']:
        section_heading("Certifications")
        pdf.set_font("Times", "", 10)
        for line in st.session_state['certifications'].split('\n'):
            line = line.strip()
            if line:
                sanitized = sanitize_text(line)
                if not sanitized.strip().startswith(chr(183)) and not sanitized.strip().startswith("-"):
                    sanitized = chr(183) + " " + sanitized
                pdf.multi_cell(0, 4, sanitized)
        pdf.ln(2)

    # --- ACHIEVEMENTS ---
    if st.session_state['achievements']:
        section_heading("Achievements")
        pdf.set_font("Times", "", 10)
        for line in st.session_state['achievements'].split('\n'):
            line = line.strip()
            if line:
                sanitized = sanitize_text(line)
                if not sanitized.strip().startswith(chr(183)) and not sanitized.strip().startswith("-"):
                    sanitized = chr(183) + " " + sanitized
                pdf.multi_cell(0, 4, sanitized)
        pdf.ln(2)

    return pdf.output(dest='S').encode('latin-1', 'replace')


# --- DOCX GENERATION LOGIC ---
def create_docx_resume():
    doc = Document()
    
    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # Heading Style
    h_style = doc.styles['Heading 1']
    h_font = h_style.font
    h_font.name = 'Times New Roman'
    h_font.size = Pt(24)
    h_font.bold = True
    h_font.color.rgb = None # Default black
    
    h2_style = doc.styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = 'Times New Roman'
    h2_font.size = Pt(14)
    h2_font.bold = True
    h2_font.small_caps = True
    h2_font.color.rgb = None
    
    # --- HEADER ---
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = header_para.add_run(name.upper() if name else "YOUR NAME")
    name_run.bold = True
    name_run.font.size = Pt(22)
    
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_info = []
    if phone: contact_info.append(phone)
    if email: contact_info.append(email)
    contact_para.add_run(" | ".join(contact_info))
    
    links_para = doc.add_paragraph()
    links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    links = []
    if linkedin: links.append(f"LinkedIn: {linkedin}")
    if github: links.append(f"GitHub: {github}")
    if hackerrank: links.append(f"HackerRank: {hackerrank}")
    if leetcode: links.append(f"LeetCode: {leetcode}")
    links_para.add_run(" | ".join(links))
    
    doc.add_paragraph() # Spacer

    # --- OBJECTIVE ---
    if st.session_state['objective']:
        doc.add_heading('OBJECTIVE', level=2)
        doc.add_paragraph(st.session_state['objective'])

    # --- EDUCATION ---
    has_education = any(e['degree'] for e in st.session_state['education_entries'])
    if has_education:
        doc.add_heading('EDUCATION', level=2)
        for edu in st.session_state['education_entries']:
            if edu['degree']:
                p = doc.add_paragraph()
                p.add_run(edu['degree']).bold = True
                p.add_run(f"\t{edu['year']}").bold = True 
                p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), 2) # Right align tab
                doc.add_paragraph(edu['school'])

    # --- SKILLS ---
    has_skills = any([skill_prog_langs, skill_web_dev, skill_core, skill_tools, skill_soft])
    if has_skills:
        doc.add_heading('TECHNICAL SKILLS', level=2)
        if skill_prog_langs: doc.add_paragraph(f"Programming Languages: {skill_prog_langs}")
        if skill_web_dev: doc.add_paragraph(f"Web Development: {skill_web_dev}")
        if skill_core: doc.add_paragraph(f"Core Concepts: {skill_core}")
        if skill_tools: doc.add_paragraph(f"Tools: {skill_tools}")
        if skill_soft: doc.add_paragraph(f"Soft Skills: {skill_soft}")

    # --- INTERNSHIPS ---
    has_internships = any(i['role'] for i in st.session_state['internship_entries'])
    if has_internships:
        doc.add_heading('INTERNSHIPS', level=2)
        for intern in st.session_state['internship_entries']:
            if intern['role']:
                p = doc.add_paragraph()
                p.add_run(intern['role']).bold = True
                p.add_run(f"\t{intern['dates']}").italic = True
                p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), 2)
                
                if intern['company']:
                    doc.add_paragraph(intern['company']).italic = True
                
                if intern['description']:
                    for line in intern['description'].split('\n'):
                        if line.strip():
                            doc.add_paragraph(line.strip(), style='List Bullet')

    # --- PROJECTS ---
    has_projects = any(p['title'] for p in st.session_state['project_entries'])
    if has_projects:
        doc.add_heading('PROJECTS', level=2)
        for proj in st.session_state['project_entries']:
            if proj['title']:
                p = doc.add_paragraph()
                p.add_run(proj['title']).bold = True
                if proj['link']:
                    p.add_run(f" ({proj['link']})").italic = True
                
                if proj['description']:
                    for line in proj['description'].split('\n'):
                        if line.strip():
                            doc.add_paragraph(line.strip(), style='List Bullet')

    # --- CERTIFICATIONS ---
    if st.session_state['certifications']:
        doc.add_heading('CERTIFICATIONS', level=2)
        for line in st.session_state['certifications'].split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip(), style='List Bullet')

    # --- ACHIEVEMENTS ---
    if st.session_state['achievements']:
        doc.add_heading('ACHIEVEMENTS', level=2)
        for line in st.session_state['achievements'].split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip(), style='List Bullet')
    
    # Save to IO
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io


# --- UI LAYOUT ---
c1, c2 = st.columns([1.2, 1])

# --- LEFT COLUMN: INPUTS ---
with c1:
    st.subheader("1. Personal Details")
    col_p1, col_p2 = st.columns(2)
    with col_p1:
        name = st.text_input("Full Name", placeholder="NALLI NEEHARIKA")
        phone = st.text_input("Phone Number", placeholder="+91 8464835326")
        linkedin = st.text_input("LinkedIn Profile", placeholder="LinkedIn URL")
        hackerrank = st.text_input("HackerRank Profile", placeholder="HackerRank URL")
    with col_p2:
        email = st.text_input("Email Address", placeholder="nallineeharika33@gmail.com")
        github = st.text_input("GitHub Profile", placeholder="GitHub URL")
        leetcode = st.text_input("LeetCode Profile", placeholder="LeetCode URL")

    st.subheader("2. Objective & Auto-Description")
    selected_role = st.selectbox("Select Resume Target Role (Auto-populates Objective)", list(JOB_DESCRIPTIONS.keys()))
    
    # Auto-populate logic: If user selects a role, update the objective in session state
    # We use a callback-like pattern here: check if role changed from a previous state if needed, 
    # but simplest is to just check if selected_role is valid and text area is empty OR user explicitly wants to overwrite.
    # To keep it non-intrusive: We only auto-fill if the user changes the selectbox.
    # We can track the 'last_selected_role' to detect change.
    
    if 'last_selected_role' not in st.session_state:
        st.session_state['last_selected_role'] = "Select a Role"
        
    if selected_role != st.session_state['last_selected_role']:
        st.session_state['last_selected_role'] = selected_role
        if selected_role != "Select a Role":
            st.session_state['objective'] = JOB_DESCRIPTIONS[selected_role]
            st.toast(f"Objective updated for {selected_role}!", icon="✨")

    st.session_state['objective'] = st.text_area("Career Objective", value=st.session_state['objective'], height=100, placeholder="Passionate about technology and coding...")

    st.subheader("3. Education")
    for i, edu in enumerate(st.session_state['education_entries']):
        st.markdown(f"**Education {i+1}**")
        c_e1, c_e2, c_e3 = st.columns([2, 2, 1])
        with c_e1:
            st.session_state['education_entries'][i]['degree'] = st.text_input(f"Degree/Course #{i+1}", value=edu['degree'], placeholder="B.Tech in CSE")
        with c_e2:
            st.session_state['education_entries'][i]['school'] = st.text_input(f"College/School #{i+1}", value=edu['school'], placeholder="College Name")
        with c_e3:
            st.session_state['education_entries'][i]['year'] = st.text_input(f"Year/Grade #{i+1}", value=edu['year'], placeholder="2022-2026")
    
    if st.button("➕ Add Education"):
        add_education()
        st.rerun()
    if len(st.session_state['education_entries']) > 1 and st.button("➖ Remove Last Education"):
        remove_education(-1)
        st.rerun()
    st.subheader("4. Technical Skills")
    st.info("Enter skills separated by commas.")
    skill_prog_langs = st.text_input("Programming Languages", placeholder="Python, Java, C++")
    skill_web_dev = st.text_input("Web Development", placeholder="HTML, CSS, JavaScript")
    skill_core = st.text_input("Core Concepts", placeholder="DSA, OOPs, DBMS, OS")
    skill_tools = st.text_input("Tools/Technologies", placeholder="VS Code, Git, Jupyter")
    skill_soft = st.text_input("Soft Skills", placeholder="Leadership, Communication")

    st.subheader("5. Internships")
    for i, intern in enumerate(st.session_state['internship_entries']):
        st.markdown(f"**Internship {i+1}**")
        st.session_state['internship_entries'][i]['role'] = st.text_input(f"Role/Title #{i+1}", value=intern['role'], placeholder="AI-ML Virtual Internship")
        c_i1, c_i2 = st.columns(2)
        with c_i1:
            st.session_state['internship_entries'][i]['company'] = st.text_input(f"Company/Org #{i+1}", value=intern['company'], placeholder="Google / AICTE")
        with c_i2:
            st.session_state['internship_entries'][i]['dates'] = st.text_input(f"Duration #{i+1}", value=intern['dates'], placeholder="May 2023 - July 2023")
        st.session_state['internship_entries'][i]['description'] = st.text_area(f"Description (Bullet points) #{i+1}", value=intern['description'], height=100, placeholder="• Gained hands-on experience...\n• Developed model for...")
    
    if st.button("➕ Add Internship"):
        add_internship()
        st.rerun()
    if len(st.session_state['internship_entries']) > 1 and st.button("➖ Remove Last Internship"):
        remove_internship(-1)
        st.rerun()

    st.subheader("6. Projects")
    for i, proj in enumerate(st.session_state['project_entries']):
        st.markdown(f"**Project {i+1}**")
        c_p1, c_p2 = st.columns([3, 1])
        with c_p1:
            st.session_state['project_entries'][i]['title'] = st.text_input(f"Project Title #{i+1}", value=proj['title'], placeholder="Attendance Automation Solution")
        with c_p2:
            st.session_state['project_entries'][i]['link'] = st.text_input(f"Link (Optional) #{i+1}", value=proj['link'], placeholder="Project URL")
        st.session_state['project_entries'][i]['description'] = st.text_area(f"Description (Bullet points) #{i+1}", value=proj['description'], height=100, placeholder="• Designed web-based system...\n• Integrated front-end and back-end...")

    if st.button("➕ Add Project"):
        add_project()
        st.rerun()
    if len(st.session_state['project_entries']) > 1 and st.button("➖ Remove Last Project"):
        remove_project(-1)
        st.rerun()

    st.subheader("7. Certifications")
    st.session_state['certifications'] = st.text_area("Certifications (One per line)", value=st.session_state['certifications'], height=100, placeholder="Salesforce Certified AI Associate\nServicenow Certified System Administrator")

    st.subheader("8. Achievements")
    st.session_state['achievements'] = st.text_area("Achievements (One per line)", value=st.session_state['achievements'], height=100, placeholder="Earned HackerRank badges\nCompleted NPTEL course")

# --- RIGHT COLUMN: LIVE PREVIEW ---
with c2:
    st.subheader("📄 Live Preview")
    
    # Download Button Top - PDF
    if name:
        col_d1, col_d2 = st.columns(2)
        with col_d1:
            try:
                pdf_bytes = create_ats_resume()
                b64_pdf = base64.b64encode(pdf_bytes).decode('latin-1')
                href_pdf = f'<a href="data:application/pdf;base64,{b64_pdf}" download="{name.replace(" ", "_")}_ATS_Resume.pdf" style="text-decoration: none;"><button style="width: 100%; background: #2c3e50; color: white; border: none; padding: 10px; border-radius: 5px; font-weight: bold; cursor: pointer; font-size: 16px; margin-bottom: 20px;">📥 PDF</button></a>'
                st.markdown(href_pdf, unsafe_allow_html=True)
            except Exception as e:
                st.error(f"Error PDF: {e}")
        
        with col_d2:
            try:
                docx_io = create_docx_resume()
                st.download_button(
                    label="📥 DOCX",
                    data=docx_io,
                    file_name=f"{name.replace(' ', '_')}_Resume.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="docx_download"
                )
            except Exception as e:
                st.error(f"Error DOCX: {e}")

    else:
        st.warning("Enter your full name to enable download.")

    # --- HTML PREVIEW CONSTRUCTION ---
    
    preview_content = ""
    
    # Header - Using textwrap.dedent to avoid Markdown Indentation issues
    preview_content += textwrap.dedent(f"""
    <div style="font-family: 'Times New Roman', Times, serif; color: black; padding: 20px; background: white; border: 1px solid #ddd; border-radius: 5px; min-height: 800px; box-shadow: 0 4px 8px rgba(0,0,0,0.1);">
        <h1 style="text-align: center; text-transform: uppercase; margin-bottom: 5px; font-size: 28px; color: #000;">{name if name else "YOUR NAME"}</h1>
        <div style="text-align: center; font-size: 14px; color: blue;">
            {phone if phone else "+91 0000000000"} | {email if email else "email@example.com"}
        </div>
        <div style="text-align: center; font-size: 14px; color: blue; margin-top: 5px;">
            {' | '.join(filter(None, [l for l in ['LinkedIn' if linkedin else '', 'GitHub' if github else '', 'HackerRank' if hackerrank else '', 'LeetCode' if leetcode else '']]))}
        </div>
        <hr style="border: 0; border-top: 1px solid #eee; margin: 15px 0;">
    """)
    
    # Helper to clean bullets for HTML
    def html_bullets(text):
        if not text: return ""
        lines = text.split('\n')
        html_list = "<ul style='margin-top: 2px; margin-bottom: 5px; padding-left: 20px;'>"
        for line in lines:
            line = line.strip()
            if line:
                # Remove inputs char bullet if present
                clean_line = line.replace('•', '').replace(chr(183), '').strip()
                if clean_line.startswith('-'): clean_line = clean_line[1:].strip()
                html_list += f"<li style='margin-bottom: 2px;'>{clean_line}</li>"
        html_list += "</ul>"
        return html_list

    # Objective
    if st.session_state['objective']:
        preview_content += textwrap.dedent(f"""
        <h3 style="text-transform: uppercase; font-size: 16px; border-bottom: 1px solid black; padding-bottom: 2px; margin-top: 20px;">Objective</h3>
        <p style="font-size: 14px; line-height: 1.4;">{st.session_state['objective']}</p>
        """)

    # Education
    if any(e['degree'] for e in st.session_state['education_entries']):
        preview_content += textwrap.dedent(f"""
        <h3 style="text-transform: uppercase; font-size: 16px; border-bottom: 1px solid black; padding-bottom: 2px; margin-top: 20px;">Education</h3>
        """)
        for edu in st.session_state['education_entries']:
            if edu['degree']:
                preview_content += textwrap.dedent(f"""
                <div style="margin-bottom: 10px; font-size: 14px;">
                    <div style="display: flex; justify-content: space-between;">
                        <strong style="width: 70%;">{edu['degree']}</strong>
                        <span style="font-weight: bold;">{edu['year']}</span>
                    </div>
                    <div>{edu['school']}</div>
                </div>
                """)

    # Skills
    if any([skill_prog_langs, skill_web_dev, skill_core, skill_tools, skill_soft]):
        preview_content += textwrap.dedent(f"""
        <h3 style="text-transform: uppercase; font-size: 16px; border-bottom: 1px solid black; padding-bottom: 2px; margin-top: 20px;">Technical Skills</h3>
        <table style="width: 100%; font-size: 14px; border-collapse: collapse;">
        """)
        
        def skill_row_html(title, val):
            if val:
                return textwrap.dedent(f"""
                <tr>
                    <td style="width: 35%; font-weight: bold; vertical-align: top; padding-bottom: 4px;">{title}</td>
                    <td style="vertical-align: top; padding-bottom: 4px;">{val}</td>
                </tr>
                """)
            return ""

        preview_content += skill_row_html("Programming Languages", skill_prog_langs)
        preview_content += skill_row_html("Web Development", skill_web_dev)
        preview_content += skill_row_html("Core Concepts", skill_core)
        preview_content += skill_row_html("Tools/Technologies", skill_tools)
        preview_content += skill_row_html("Soft Skills", skill_soft)
        preview_content += "</table>"

    # Internships
    if any(i['role'] for i in st.session_state['internship_entries']):
        preview_content += textwrap.dedent(f"""
        <h3 style="text-transform: uppercase; font-size: 16px; border-bottom: 1px solid black; padding-bottom: 2px; margin-top: 20px;">Internships</h3>
        """)
        for intern in st.session_state['internship_entries']:
            if intern['role']:
                preview_content += textwrap.dedent(f"""
                <div style="margin-bottom: 10px; font-size: 14px;">
                    <div style="font-weight: bold;">{intern['role']}</div>
                    <div style="display: flex; justify-content: space-between; font-style: italic;">
                        <span>{intern['company']}</span>
                        <span>{intern['dates']}</span>
                    </div>
                    {html_bullets(intern['description'])}
                </div>
                """)

    # Projects
    if any(p['title'] for p in st.session_state['project_entries']):
        preview_content += textwrap.dedent(f"""
        <h3 style="text-transform: uppercase; font-size: 16px; border-bottom: 1px solid black; padding-bottom: 2px; margin-top: 20px;">Projects</h3>
        """)
        for proj in st.session_state['project_entries']:
            if proj['title']:
                link_html = f'<a href="{proj["link"]}" style="color: blue; text-decoration: none;">Project Link</a>' if proj['link'] else ''
                preview_content += textwrap.dedent(f"""
                <div style="margin-bottom: 10px; font-size: 14px;">
                    <div style="display: flex; justify-content: space-between;">
                        <strong style="width: 75%;">{proj['title']}</strong>
                        <span style="font-size: 12px;">{link_html}</span>
                    </div>
                    {html_bullets(proj['description'])}
                </div>
                """)

    # Certifications
    if st.session_state['certifications']:
        preview_content += textwrap.dedent(f"""
        <h3 style="text-transform: uppercase; font-size: 16px; border-bottom: 1px solid black; padding-bottom: 2px; margin-top: 20px;">Certifications</h3>
        <div style="font-size: 14px;">{html_bullets(st.session_state['certifications'])}</div>
        """)

    # Achievements
    if st.session_state['achievements']:
        preview_content += textwrap.dedent(f"""
        <h3 style="text-transform: uppercase; font-size: 16px; border-bottom: 1px solid black; padding-bottom: 2px; margin-top: 20px;">Achievements</h3>
            <div style="font-size: 14px;">{html_bullets(st.session_state['achievements'])}</div>
        """)
        
    preview_content += "</div>"
    
    # Render
    components.html(preview_content, height=1000, scrolling=True)
